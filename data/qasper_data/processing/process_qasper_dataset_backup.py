#!/usr/bin/env python3
"""
Script backup để xử lý QASPER Dataset với dữ liệu mẫu
Nếu không download được từ Hugging Face, sử dụng dữ liệu mẫu để test
"""

import os
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import Dict, List, Any
from pathlib import Path

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class QASPERProcessorBackup:
    def __init__(self, output_dir: str = "data/qasper_processed"):
        self.output_dir = Path(output_dir)
        self.docx_dir = self.output_dir / "docx_files"
        self.qa_dir = self.output_dir / "qa_data"
        
        # Tạo thư mục output
        self.docx_dir.mkdir(parents=True, exist_ok=True)
        self.qa_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Output directories created: {self.output_dir}")
    
    def create_sample_data(self) -> List[Dict]:
        """Tạo dữ liệu mẫu để test"""
        logger.info("Creating sample data for testing...")
        
        sample_records = []
        
        # Tạo 5 bản ghi mẫu
        for i in range(5):
            record = {
                'id': f'sample_{i+1:03d}',
                'title': f'Sample Research Paper {i+1}: Advanced Machine Learning Techniques',
                'abstract': f'This is a sample abstract for paper {i+1}. It describes the methodology and results of our research on advanced machine learning techniques for natural language processing.',
                'full_text': {
                    'section_name': [
                        'Introduction',
                        'Related Work', 
                        'Methodology',
                        'Experiments',
                        'Results',
                        'Conclusion'
                    ],
                    'paragraphs': [
                        [f'This is the introduction section of sample paper {i+1}.'],
                        [f'This section reviews related work in the field for paper {i+1}.'],
                        [f'Our methodology for paper {i+1} is described here.'],
                        [f'Experimental setup and results for paper {i+1}.'],
                        [f'Analysis of results for paper {i+1}.'],
                        [f'Conclusion and future work for paper {i+1}.']
                    ]
                },
                'qas': {
                    'question': [
                        f'What is the main contribution of paper {i+1}?',
                        f'What methodology is used in paper {i+1}?',
                        f'What are the results of paper {i+1}?'
                    ],
                    'question_id': [
                        f'q1_paper_{i+1}',
                        f'q2_paper_{i+1}',
                        f'q3_paper_{i+1}'
                    ],
                    'nlp_background': ['two', 'two', 'zero'],
                    'topic_background': ['unfamiliar', 'unfamiliar', 'unfamiliar'],
                    'paper_read': ['no', 'no', 'no'],
                    'search_query': ['', '', ''],
                    'question_writer': [
                        f'writer_{i+1}',
                        f'writer_{i+1}',
                        f'writer_{i+1}'
                    ],
                    'answers': [
                        {
                            'answer': [
                                {
                                    'unanswerable': False,
                                    'extractive_spans': [],
                                    'yes_no': None,
                                    'free_form_answer': f'The main contribution of paper {i+1} is the development of advanced ML techniques.',
                                    'evidence': [f'Paper {i+1} presents novel machine learning approaches.'],
                                    'highlighted_evidence': [f'Paper {i+1} presents novel machine learning approaches.']
                                }
                            ],
                            'annotation_id': [f'anno_{i+1}_1'],
                            'worker_id': [f'worker_{i+1}']
                        },
                        {
                            'answer': [
                                {
                                    'unanswerable': False,
                                    'extractive_spans': [],
                                    'yes_no': None,
                                    'free_form_answer': f'Paper {i+1} uses deep learning methodology.',
                                    'evidence': [f'The methodology section describes deep learning approaches.'],
                                    'highlighted_evidence': [f'The methodology section describes deep learning approaches.']
                                }
                            ],
                            'annotation_id': [f'anno_{i+1}_2'],
                            'worker_id': [f'worker_{i+1}']
                        },
                        {
                            'answer': [
                                {
                                    'unanswerable': False,
                                    'extractive_spans': [],
                                    'yes_no': None,
                                    'free_form_answer': f'Paper {i+1} achieves 95% accuracy.',
                                    'evidence': [f'The results show 95% accuracy on the test set.'],
                                    'highlighted_evidence': [f'The results show 95% accuracy on the test set.']
                                }
                            ],
                            'annotation_id': [f'anno_{i+1}_3'],
                            'worker_id': [f'worker_{i+1}']
                        }
                    ]
                }
            }
            sample_records.append(record)
        
        logger.info(f"Created {len(sample_records)} sample records")
        return sample_records
    
    def create_docx_from_record(self, record: Dict) -> str:
        """Tạo file .docx từ một bản ghi QASPER"""
        doc_id = record['id']
        title = record['title']
        abstract = record['abstract']
        full_text = record['full_text']
        
        # Tạo document mới
        doc = Document()
        
        # Thêm title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm abstract
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(abstract)
        
        # Thêm full text
        doc.add_heading('Full Text', level=1)
        
        # Xử lý full_text structure
        if 'section_name' in full_text and 'paragraphs' in full_text:
            sections = full_text['section_name']
            paragraphs = full_text['paragraphs']
            
            for i, section in enumerate(sections):
                # Thêm section heading
                doc.add_heading(section, level=2)
                
                # Thêm paragraphs cho section này
                if i < len(paragraphs):
                    section_paragraphs = paragraphs[i]
                    for para_text in section_paragraphs:
                        if para_text.strip():  # Chỉ thêm paragraph không rỗng
                            doc.add_paragraph(para_text)
        else:
            # Fallback: thêm full_text như plain text
            doc.add_paragraph(str(full_text))
        
        # Lưu file
        docx_filename = f"{doc_id}.docx"
        docx_path = self.docx_dir / docx_filename
        doc.save(docx_path)
        
        logger.info(f"Created DOCX: {docx_filename}")
        return docx_filename
    
    def extract_qa_data(self, record: Dict) -> Dict:
        """Trích xuất và chuẩn hóa QA data theo format simple_qa_dataset.json"""
        qas = record.get('qas', {})
        
        if not qas:
            return None
        
        # Chuẩn hóa theo format bạn yêu cầu
        qa_data = {
            "question": qas.get('question', []),
            "question_id": qas.get('question_id', []),
            "nlp_background": qas.get('nlp_background', []),
            "topic_background": qas.get('topic_background', []),
            "paper_read": qas.get('paper_read', []),
            "search_query": qas.get('search_query', []),
            "question_writer": qas.get('question_writer', []),
            "answers": qas.get('answers', [])
        }
        
        return qa_data
    
    def process_dataset(self, limit: int = 5):
        """Xử lý dataset với dữ liệu mẫu"""
        logger.info(f"Starting QASPER dataset processing with sample data (limit: {limit} records)...")
        
        # Tạo dữ liệu mẫu
        dataset = self.create_sample_data()
        
        processed_count = 0
        qa_records = []
        
        for record in dataset:
            if processed_count >= limit:
                logger.info(f"Reached limit of {limit} records, stopping...")
                break
            try:
                # Tạo DOCX file
                docx_filename = self.create_docx_from_record(record)
                
                # Trích xuất QA data
                qa_data = self.extract_qa_data(record)
                if qa_data:
                    qa_data['source_document_id'] = record['id']
                    qa_data['docx_filename'] = docx_filename
                    qa_records.append(qa_data)
                
                processed_count += 1
                
                if processed_count % 2 == 0:
                    logger.info(f"Processed {processed_count}/{limit} records")
                    
            except Exception as e:
                logger.error(f"Error processing record {record.get('id', 'unknown')}: {e}")
                continue
        
        # Lưu QA data
        qa_output_file = self.qa_dir / "qasper_qa_dataset.json"
        with open(qa_output_file, 'w', encoding='utf-8') as f:
            json.dump(qa_records, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Processing completed!")
        logger.info(f"- Processed {processed_count}/{limit} records")
        logger.info(f"- Created {len(os.listdir(self.docx_dir))} DOCX files")
        logger.info(f"- Saved QA data to {qa_output_file}")
        logger.info(f"- DOCX files saved in: {self.docx_dir}")
        logger.info(f"- QA data saved in: {self.qa_dir}")

def main():
    """Main function"""
    processor = QASPERProcessorBackup()
    processor.process_dataset(limit=5)  # Test với 5 tài liệu mẫu

if __name__ == "__main__":
    main()
