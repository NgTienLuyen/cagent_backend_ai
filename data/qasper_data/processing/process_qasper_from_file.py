#!/usr/bin/env python3
"""
Script để xử lý QASPER dataset từ file đã download
"""

import os
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import Dict, List, Any
from pathlib import Path

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class QASPERProcessorFromFile:
    def __init__(self, output_dir: str = "data/qasper_processed"):
        self.output_dir = Path(output_dir)
        self.docx_dir = self.output_dir / "docx_files"
        self.qa_dir = self.output_dir / "qa_data"
        
        # Tạo thư mục output
        self.docx_dir.mkdir(parents=True, exist_ok=True)
        self.qa_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Output directories created: {self.output_dir}")
    
    def load_qasper_data(self, file_path: str = "data/qasper_raw/validation.json") -> List[Dict]:
        """Load dữ liệu QASPER từ file"""
        logger.info(f"Loading QASPER data from: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            logger.info(f"Loaded {len(data)} records from file")
            return data
        except Exception as e:
            logger.error(f"Error loading data from file: {e}")
            raise
    
    def create_docx_from_record(self, record: Dict) -> str:
        """Tạo file .docx từ một bản ghi QASPER"""
        doc_id = record['id']
        title = record['title']
        abstract = record['abstract']
        full_text = record['full_text']
        
        # Tạo document mới
        doc = Document()
        
        # Thêm title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm abstract
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(abstract)
        
        # Thêm full text
        doc.add_heading('Full Text', level=1)
        
        # Xử lý full_text structure
        if 'section_name' in full_text and 'paragraphs' in full_text:
            sections = full_text['section_name']
            paragraphs = full_text['paragraphs']
            
            for i, section in enumerate(sections):
                # Thêm section heading
                doc.add_heading(section, level=2)
                
                # Thêm paragraphs cho section này
                if i < len(paragraphs):
                    section_paragraphs = paragraphs[i]
                    for para_text in section_paragraphs:
                        if para_text.strip():  # Chỉ thêm paragraph không rỗng
                            doc.add_paragraph(para_text)
        else:
            # Fallback: thêm full_text như plain text
            doc.add_paragraph(str(full_text))
        
        # Lưu file
        docx_filename = f"{doc_id}.docx"
        docx_path = self.docx_dir / docx_filename
        doc.save(docx_path)
        
        logger.info(f"Created DOCX: {docx_filename}")
        return docx_filename
    
    def extract_qa_data(self, record: Dict) -> List[Dict]:
        """Trích xuất và chuẩn hóa QA data theo format simple_qa_dataset.json"""
        qas = record.get('qas', {})
        
        if not qas:
            return []
        
        # Convert từ QASPER format sang simple_qa_dataset format
        questions = qas.get('question', [])
        answers = qas.get('answers', [])
        
        qa_records = []
        
        for i, question in enumerate(questions):
            if i < len(answers):
                answer_data = answers[i]
                
                # Lấy câu trả lời tốt nhất
                ground_truth_answer = ""
                relevant_contexts = []
                
                if answer_data and 'answer' in answer_data:
                    answer_list = answer_data['answer']
                    if answer_list:
                        best_answer = answer_list[0]  # Lấy answer đầu tiên
                        
                        # Ưu tiên free_form_answer, sau đó extractive_spans
                        if best_answer.get('free_form_answer'):
                            ground_truth_answer = best_answer['free_form_answer']
                        elif best_answer.get('extractive_spans'):
                            ground_truth_answer = " ".join(best_answer['extractive_spans'])
                        
                        # Lấy evidence làm relevant_contexts
                        if best_answer.get('evidence'):
                            relevant_contexts = best_answer['evidence']
                
                # Tạo record theo format simple_qa_dataset
                qa_record = {
                    "question": question,
                    "ground_truth_answer": ground_truth_answer,
                    "relevant_contexts": relevant_contexts,
                    "relevant_chunk_ids": [f"qasper_{record['id']}_{i}"],  # Tạo fake chunk_id
                    "source_document_id": record['id']
                }
                
                qa_records.append(qa_record)
        
        return qa_records
    
    def process_dataset(self, limit: int = 20, input_file: str = None):
        """Xử lý dataset từ file"""
        logger.info(f"Starting QASPER dataset processing from file (limit: {limit} records)...")
        
        # Tự động tìm file validation.json
        if input_file is None:
            possible_paths = [
                "data/data/qasper_raw/validation.json",
                "../data/data/qasper_raw/validation.json", 
                "data/qasper_raw/validation.json",
                "../data/qasper_raw/validation.json"
            ]
            
            for path in possible_paths:
                if os.path.exists(path):
                    input_file = path
                    logger.info(f"Found validation file at: {input_file}")
                    break
            
            if input_file is None:
                raise FileNotFoundError("Could not find validation.json file")
        
        # Load dữ liệu từ file
        dataset = self.load_qasper_data(input_file)
        
        processed_count = 0
        qa_records = []
        
        for record in dataset:
            if processed_count >= limit:
                logger.info(f"Reached limit of {limit} records, stopping...")
                break
            try:
                # Tạo DOCX file
                docx_filename = self.create_docx_from_record(record)
                
                # Trích xuất QA data (trả về list các QA records)
                qa_data_list = self.extract_qa_data(record)
                if qa_data_list:
                    # Thêm từng QA record vào danh sách
                    for qa_record in qa_data_list:
                        qa_record['docx_filename'] = docx_filename
                        qa_records.append(qa_record)
                
                processed_count += 1
                
                if processed_count % 5 == 0:
                    logger.info(f"Processed {processed_count}/{limit} records")
                    
            except Exception as e:
                logger.error(f"Error processing record {record.get('id', 'unknown')}: {e}")
                continue
        
        # Lưu QA data
        qa_output_file = self.qa_dir / "qasper_qa_dataset.json"
        with open(qa_output_file, 'w', encoding='utf-8') as f:
            json.dump(qa_records, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Processing completed!")
        logger.info(f"- Processed {processed_count}/{limit} records")
        logger.info(f"- Created {len(os.listdir(self.docx_dir))} DOCX files")
        logger.info(f"- Saved QA data to {qa_output_file}")
        logger.info(f"- DOCX files saved in: {self.docx_dir}")
        logger.info(f"- QA data saved in: {self.qa_dir}")

def main():
    """Main function"""
    processor = QASPERProcessorFromFile()
    processor.process_dataset(limit=20)  # Xử lý 20 tài liệu đầu tiên

if __name__ == "__main__":
    main()
