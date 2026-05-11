#!/usr/bin/env python3
"""
Script để xử lý QASPER Dataset từ Hugging Face
- Download validation set từ allenai/qasper
- Tách title + abstract + full_text thành file .docx
- Chuẩn hóa cấu trúc QA theo format simple_qa_dataset.json
"""

import os
import json
import requests
from datasets import load_dataset
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import Dict, List, Any
from pathlib import Path

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class QASPERProcessor:
    def __init__(self, output_dir: str = "data/qasper_processed"):
        self.output_dir = Path(output_dir)
        self.docx_dir = self.output_dir / "docx_files"
        self.qa_dir = self.output_dir / "qa_data"
        
        # Tạo thư mục output
        self.docx_dir.mkdir(parents=True, exist_ok=True)
        self.qa_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Output directories created: {self.output_dir}")
    
    def download_qasper_validation(self) -> List[Dict]:
        """Download validation set từ QASPER dataset"""
        logger.info("Downloading QASPER validation dataset...")
        
        try:
            # Thử cách 1: Load với ignore_verifications=True
            dataset = load_dataset("allenai/qasper", split="validation", ignore_verifications=True)
            logger.info(f"Downloaded {len(dataset)} validation records")
            return dataset
        except Exception as e:
            logger.error(f"Method 1 failed: {e}")
            logger.info("Trying method 2...")
            
            try:
                # Thử cách 2: Load toàn bộ dataset rồi lấy validation
                full_dataset = load_dataset("allenai/qasper", ignore_verifications=True)
                validation_dataset = full_dataset["validation"]
                logger.info(f"Downloaded {len(validation_dataset)} validation records (method 2)")
                return validation_dataset
            except Exception as e2:
                logger.error(f"Method 2 failed: {e2}")
                logger.info("Trying method 3...")
                
                try:
                    # Thử cách 3: Load với download_mode="reuse_dataset_if_exists"
                    dataset = load_dataset("allenai/qasper", split="validation", download_mode="reuse_dataset_if_exists")
                    logger.info(f"Downloaded {len(dataset)} validation records (method 3)")
                    return dataset
                except Exception as e3:
                    logger.error(f"Method 3 failed: {e3}")
                    logger.info("Trying method 4...")
                    
                    try:
                        # Thử cách 4: Load từ cache hoặc download manual
                        from datasets import Dataset
                        import requests
                        
                        # Download manual từ Hugging Face Hub
                        logger.info("Attempting manual download...")
                        dataset = load_dataset("allenai/qasper", split="validation", streaming=True)
                        
                        # Convert streaming dataset to list
                        records = []
                        for i, record in enumerate(dataset):
                            records.append(record)
                            if i >= 100:  # Limit để test
                                break
                        
                        logger.info(f"Downloaded {len(records)} validation records (manual method)")
                        return records
                        
                    except Exception as e4:
                        logger.error(f"All methods failed. Last error: {e4}")
                        raise Exception("Could not download QASPER dataset. Please check your internet connection and try again.")
    
    def create_docx_from_record(self, record: Dict) -> str:
        """Tạo file .docx từ một bản ghi QASPER"""
        doc_id = record['id']
        title = record['title']
        abstract = record['abstract']
        full_text = record['full_text']
        
        # Tạo document mới
        doc = Document()
        
        # Thêm title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm abstract
        doc.add_heading('Abstract', level=1)
        doc.add_paragraph(abstract)
        
        # Thêm full text
        doc.add_heading('Full Text', level=1)
        
        # Xử lý full_text structure
        if 'section_name' in full_text and 'paragraphs' in full_text:
            sections = full_text['section_name']
            paragraphs = full_text['paragraphs']
            
            for i, section in enumerate(sections):
                # Thêm section heading
                doc.add_heading(section, level=2)
                
                # Thêm paragraphs cho section này
                if i < len(paragraphs):
                    section_paragraphs = paragraphs[i]
                    for para_text in section_paragraphs:
                        if para_text.strip():  # Chỉ thêm paragraph không rỗng
                            doc.add_paragraph(para_text)
        else:
            # Fallback: thêm full_text như plain text
            doc.add_paragraph(str(full_text))
        
        # Lưu file
        docx_filename = f"{doc_id}.docx"
        docx_path = self.docx_dir / docx_filename
        doc.save(docx_path)
        
        logger.info(f"Created DOCX: {docx_filename}")
        return docx_filename
    
    def extract_qa_data(self, record: Dict) -> Dict:
        """Trích xuất và chuẩn hóa QA data theo format simple_qa_dataset.json"""
        qas = record.get('qas', {})
        
        if not qas:
            return None
        
        # Chuẩn hóa theo format bạn yêu cầu
        qa_data = {
            "question": qas.get('question', []),
            "question_id": qas.get('question_id', []),
            "nlp_background": qas.get('nlp_background', []),
            "topic_background": qas.get('topic_background', []),
            "paper_read": qas.get('paper_read', []),
            "search_query": qas.get('search_query', []),
            "question_writer": qas.get('question_writer', []),
            "answers": qas.get('answers', [])
        }
        
        return qa_data
    
    def process_dataset(self, limit: int = 20):
        """Xử lý dataset với giới hạn số lượng"""
        logger.info(f"Starting QASPER dataset processing (limit: {limit} records)...")
        
        # Download dataset
        dataset = self.download_qasper_validation()
        
        processed_count = 0
        qa_records = []
        
        for record in dataset:
            if processed_count >= limit:
                logger.info(f"Reached limit of {limit} records, stopping...")
                break
            try:
                # Tạo DOCX file
                docx_filename = self.create_docx_from_record(record)
                
                # Trích xuất QA data
                qa_data = self.extract_qa_data(record)
                if qa_data:
                    qa_data['source_document_id'] = record['id']
                    qa_data['docx_filename'] = docx_filename
                    qa_records.append(qa_data)
                
                processed_count += 1
                
                if processed_count % 5 == 0:
                    logger.info(f"Processed {processed_count}/{limit} records")
                    
            except Exception as e:
                logger.error(f"Error processing record {record.get('id', 'unknown')}: {e}")
                continue
        
        # Lưu QA data
        qa_output_file = self.qa_dir / "qasper_qa_dataset.json"
        with open(qa_output_file, 'w', encoding='utf-8') as f:
            json.dump(qa_records, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Processing completed!")
        logger.info(f"- Processed {processed_count}/{limit} records")
        logger.info(f"- Created {len(os.listdir(self.docx_dir))} DOCX files")
        logger.info(f"- Saved QA data to {qa_output_file}")
        logger.info(f"- DOCX files saved in: {self.docx_dir}")
        logger.info(f"- QA data saved in: {self.qa_dir}")

def main():
    """Main function"""
    processor = QASPERProcessor()
    processor.process_dataset(limit=20)  # Chỉ xử lý 20 tài liệu đầu tiên

if __name__ == "__main__":
    main()
