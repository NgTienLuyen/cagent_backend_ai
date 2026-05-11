import json
import os
from docx import Document
from docx.shared import Inches
import uuid
import time
from datasets import load_dataset
from collections import defaultdict
import re

def convert_dataset_to_word_files(dataset, output_dir, dataset_name="custom_dataset"):
    """
    Chuyển đổi dataset từ load_dataset thành các file Word riêng biệt
    Nhóm theo document.id để tránh tạo file trùng lặp
    
    Args:
        dataset: Dataset object từ load_dataset
        output_dir: Thư mục output
        dataset_name: Tên dataset để đặt tên file
    """
    # Tạo thư mục output
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"✅ Đã load dataset với {len(dataset)} items")
    
    # Nhóm items theo document.id
    print("🔄 Đang nhóm items theo document ID...")
    documents_grouped = defaultdict(list)
    
    for i, item in enumerate(dataset):
        # Xác định document ID
        doc_id = None
        if 'document' in item and isinstance(item['document'], dict) and 'id' in item['document']:
            doc_id = item['document']['id']
        elif 'id' in item:
            doc_id = item['id']
        else:
            # Nếu không có ID, sử dụng index
            doc_id = f"item_{i+1}"
        
        documents_grouped[doc_id].append({
            'index': i + 1,
            'item': item
        })
    
    print(f"📊 Đã nhóm thành {len(documents_grouped)} tài liệu duy nhất")
    
    # Xử lý tất cả các document
    all_docs = list(documents_grouped.items())
    
    print(f"📊 Sẽ xử lý tất cả {len(all_docs)} tài liệu")
    print("🔄 Thứ tự xử lý (theo thứ tự xuất hiện trong dataset):")
    for i, (doc_id, items) in enumerate(all_docs):
        first_item = items[0]['item']
        kind = get_document_kind(first_item)
        print(f"   {i+1}. {doc_id[:20]}... ({kind}) - {len(items)} Q&A pairs")
    
    converted_files = []
    
    for doc_id, items in all_docs:
        print(f"\n📝 Đang xử lý tài liệu {doc_id} với {len(items)} Q&A pairs...")
        
        try:
            # Lấy item đầu tiên để lấy thông tin tài liệu
            first_item = items[0]['item']
            
            # Tạo document Word
            doc = Document()
            
            # Lấy nội dung chính của tài liệu
            content_text, content_source = get_best_content_text(first_item)
            
            if content_text:
                print(f"   📖 Sử dụng {content_source} (độ dài: {len(content_text)} chars)")
                
                # Chia text thành các đoạn để dễ đọc
                paragraphs = content_text.split('\n')
                for para in paragraphs:
                    if para.strip():  # Chỉ thêm đoạn không rỗng
                        doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph("(Không có nội dung)")
                print(f"   ❌ Không tìm thấy nội dung trong tài liệu")
            
            # Tạo tên file với ID đầy đủ
            # Giữ nguyên ID gốc, chỉ thay thế ký tự không hợp lệ cho tên file
            safe_id = str(doc_id).replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
            filename = f"doc_{safe_id}.docx"
            filepath = os.path.join(output_dir, filename)
            
            # Lưu file
            doc.save(filepath)
            
            # Lưu thông tin về file đã chuyển đổi
            converted_files.append({
                'document_id': doc_id,
                'word_file': filename,
                'filepath': filepath,
                'qa_pairs_count': len(items),
                'original_indices': [item_data['index'] for item_data in items],
                'questions': [get_question_text(item_data['item']) for item_data in items],
                'answers': [get_answers_text(item_data['item']) for item_data in items],
                'document_kind': get_document_kind(first_item),
                'document_word_count': get_word_count(first_item, content_text),
                'content_source': content_source,
                'file_size_bytes': os.path.getsize(filepath)
            })
            
            print(f"✅ Đã tạo: {filename}")
            print(f"   📄 Document ID: {doc_id}")
            print(f"   🎬 Kind: {get_document_kind(first_item)}")
            print(f"   ❓ Q&A pairs: {len(items)}")
            print(f"   📖 Content source: {content_source}")
            
        except Exception as e:
            print(f"❌ Lỗi khi xử lý document {doc_id}: {e}")
            continue
    
    # Lưu mapping file
    mapping_file = os.path.join(output_dir, "file_mapping.json")
    with open(mapping_file, "w", encoding="utf-8") as f:
        json.dump(converted_files, f, ensure_ascii=False, indent=2)
    
    # Tạo file README
    readme_file = os.path.join(output_dir, "README.md")
    with open(readme_file, "w", encoding="utf-8") as f:
        f.write(f"""# Converted Word Files (All Documents)

## Tổng quan
Đây là các file Word được tạo tự động từ **tất cả các tài liệu** của dataset: {dataset_name}

## Thống kê
- **Tổng số file:** {len(converted_files)}
- **Tổng số Q&A pairs:** {sum(f['qa_pairs_count'] for f in converted_files)}
- **Số tài liệu được xử lý:** {len(all_docs)}
- **Thời gian tạo:** {time.strftime('%Y-%m-%d %H:%M:%S')}
- **Dataset gốc:** {dataset_name}

## Cách sử dụng
1. Upload các file Word này vào hệ thống RAG của bạn
2. Thực hiện chunking và embedding
3. Sử dụng file `file_mapping.json` để map câu hỏi với file tương ứng
4. Chạy đánh giá RAG system

## Cấu trúc file
Mỗi file Word chứa:
- Chỉ nội dung tài liệu gốc (thuần túy)
- Không có metadata, header, footer
- Không có Q&A pairs
- **Tất cả các tài liệu** được xử lý

## File mapping
Xem `file_mapping.json` để biết chi tiết về mỗi file.
""")
    
    print(f"\n🎉 Hoàn thành! Đã tạo {len(converted_files)} file Word (tất cả các tài liệu)")
    print(f"📁 Output directory: {output_dir}")
    print(f"🗂️ Mapping file: {mapping_file}")
    print(f"📖 README file: {readme_file}")
    
    # In thống kê
    if converted_files:
        total_size = sum(f['file_size_bytes'] for f in converted_files)
        avg_size = total_size / len(converted_files)
        total_qa_pairs = sum(f['qa_pairs_count'] for f in converted_files)
        
        print(f"\n📊 Thống kê (tất cả các tài liệu):")
        print(f"   📁 Tổng số file Word: {len(converted_files)}")
        print(f"   ❓ Tổng số Q&A pairs: {total_qa_pairs}")
        print(f"   📊 Trung bình Q&A pairs/file: {total_qa_pairs/len(converted_files):.1f}")
        print(f"   💾 Tổng dung lượng: {total_size:,} bytes ({total_size/1024/1024:.2f} MB)")
        print(f"   📏 Dung lượng trung bình: {avg_size:,.0f} bytes")
        print(f"   🎬 Document kinds: {', '.join(set(f['document_kind'] for f in converted_files))}")
        print(f"   📝 Mỗi file chứa: Chỉ nội dung tài liệu gốc (không có metadata)")
        print(f"   🎯 Chỉ xử lý: Tất cả các tài liệu từ tổng số {len(documents_grouped)} tài liệu")
        
        print(f"\n📋 Chi tiết từng tài liệu:")
        for i, file_info in enumerate(converted_files):
            print(f"   {i+1}. {file_info['word_file']}")
            print(f"      📄 ID: {file_info['document_id']}")
            print(f"      🎬 Kind: {file_info['document_kind']}")
            print(f"      ❓ Q&A pairs: {file_info['qa_pairs_count']}")
            print(f"      📖 Content source: {file_info['content_source']}")
            print(f"      📏 File size: {file_info['file_size_bytes']:,} bytes")
    
    return converted_files

def get_question_text(item):
    """Lấy text của câu hỏi từ item"""
    if 'question' in item:
        if isinstance(item['question'], dict) and 'text' in item['question']:
            return item['question']['text']
        else:
            return str(item['question'])
    elif 'question_text' in item:
        return item['question_text']
    return "N/A"

def get_answers_text(item):
    """Lấy text của các câu trả lời từ item"""
    answers = []
    if 'answers' in item:
        if isinstance(item['answers'], list):
            answers = item['answers']
        else:
            answers = [item['answers']]
    elif 'answer' in item:
        answers = [item['answer']]
    elif 'answer_text' in item:
        answers = [item['answer_text']]
    
    return [str(ans) if not isinstance(ans, dict) else ans.get('text', str(ans)) for ans in answers]

def get_document_kind(item):
    """Lấy kind của document từ item"""
    if 'document' in item and isinstance(item['document'], dict) and 'kind' in item['document']:
        return item['document']['kind']
    elif 'kind' in item:
        return item['kind']
    return 'unknown'

def get_word_count(item, content_text):
    """Lấy word count từ item hoặc tính từ content"""
    if 'document' in item and isinstance(item['document'], dict) and 'word_count' in item['document']:
        return item['document']['word_count']
    elif 'word_count' in item:
        return item['word_count']
    elif content_text:
        return len(content_text.split())
    return 0

def clean_html_text(html_text):
    """Làm sạch HTML text để lấy nội dung thuần túy"""
    if not html_text:
        return ""
    
    # Loại bỏ các tag HTML cơ bản
    clean_text = re.sub(r'<[^>]+>', '', html_text)
    
    # Loại bỏ các ký tự đặc biệt HTML
    clean_text = clean_text.replace('&nbsp;', ' ')
    clean_text = clean_text.replace('&amp;', '&')
    clean_text = clean_text.replace('&lt;', '<')
    clean_text = clean_text.replace('&gt;', '>')
    clean_text = clean_text.replace('&quot;', '"')
    
    # Loại bỏ khoảng trắng thừa
    clean_text = re.sub(r'\s+', ' ', clean_text)
    clean_text = clean_text.strip()
    
    return clean_text

def get_best_content_text(item):
    """Lấy nội dung tốt nhất từ item, ưu tiên summary.text"""
    if 'document' in item and isinstance(item['document'], dict):
        doc_info = item['document']
        
        # Ưu tiên 1: summary.text (nội dung tóm tắt/thực sự)
        if 'summary' in doc_info and isinstance(doc_info['summary'], dict) and 'text' in doc_info['summary']:
            return doc_info['summary']['text'], "summary.text"
        
        # Ưu tiên 2: full_text nếu có
        elif 'full_text' in doc_info:
            return doc_info['full_text'], "full_text"
        
        # Ưu tiên 3: text (có thể là HTML)
        elif 'text' in doc_info:
            text_content = doc_info['text']
            # Kiểm tra xem có phải HTML không
            if text_content and ('<' in text_content and '>' in text_content):
                # Nếu là HTML, thử làm sạch
                clean_content = clean_html_text(text_content)
                if clean_content and len(clean_content) > 100:  # Đảm bảo có nội dung sau khi làm sạch
                    return clean_content, "text (cleaned HTML)"
                else:
                    return None, "text (HTML, too short after cleaning)"
            else:
                return text_content, "text"
        
        # Ưu tiên 4: story
        elif 'story' in doc_info:
            return doc_info['story'], "story"
        
        # Ưu tiên 5: context
        elif 'context' in doc_info:
            return doc_info['context'], "context"
    
    # Fallback: kiểm tra các trường khác
    if 'text' in item:
        return item['text'], "text (fallback)"
    elif 'story' in item:
        return item['story'], "story (fallback)"
    elif 'context' in item:
        return item['context'], "context (fallback)"
    
    return None, "none"

def main():
    """Hàm chính để chạy script"""
    print("🚀 SCRIPT CHUYỂN ĐỔI DATASET THÀNH FILE WORD THUẦN TÚY")
    print("=" * 70)
    print("📝 Chỉ chứa nội dung tài liệu gốc, không có metadata")
    print("=" * 70)
    
    print("📋 Các tùy chọn:")
    print("1. Load dataset từ Hugging Face datasets")
    print("2. Load dataset từ file JSON")
    print("3. Sử dụng dataset có sẵn trong code")
    
    choice = input("\n🎯 Chọn tùy chọn (1/2/3): ").strip()
    
    if choice == "1":
        # Load từ Hugging Face datasets
        dataset_name = input("📚 Nhập tên dataset (ví dụ: deepmind/narrativeqa): ").strip()
        split_name = input("📊 Nhập split (ví dụ: validation, train, test): ").strip()
        
        try:
            print(f"🔄 Đang load dataset {dataset_name} với split {split_name}...")
            dataset = load_dataset(dataset_name, split=split_name)
            print(f"✅ Đã load thành công dataset với {len(dataset)} items")
        except Exception as e:
            print(f"❌ Lỗi khi load dataset: {e}")
            return
    
    elif choice == "2":
        # Load từ file JSON
        dataset_path = input("📁 Nhập đường dẫn đến file dataset JSON: ").strip()
        
        if not os.path.exists(dataset_path):
            print(f"❌ Không tìm thấy file dataset: {dataset_path}")
            return
        
        try:
            with open(dataset_path, 'r', encoding='utf-8') as f:
                dataset = json.load(f)
            print(f"✅ Đã load dataset từ file JSON với {len(dataset)} items")
        except Exception as e:
            print(f"❌ Lỗi khi load file JSON: {e}")
            return
    
    elif choice == "3":
        # Sử dụng dataset có sẵn
        print("🔄 Đang load dataset có sẵn...")
        dataset = load_dataset("deepmind/narrativeqa", split="validation")
        print(f"✅ Đã load dataset NarrativeQA validation với {len(dataset)} items")
    
    else:
        print("❌ Lựa chọn không hợp lệ")
        return
    
    # Nhập thông tin output
    output_dir = input("📂 Nhập thư mục output (mặc định: converted_word_files): ").strip()
    if not output_dir:
        output_dir = "converted_word_files"
    
    dataset_name = input("📝 Nhập tên dataset để hiển thị (mặc định: custom_dataset): ").strip()
    if not dataset_name:
        dataset_name = "custom_dataset"
    
    # Chạy chuyển đổi
    try:
        files = convert_dataset_to_word_files(dataset, output_dir, dataset_name)
        print(f"\n🎯 Chuyển đổi thành công! Đã tạo {len(files)} file Word (nhóm theo document ID).")
        
    except Exception as e:
        print(f"❌ Lỗi trong quá trình chuyển đổi: {e}")

if __name__ == "__main__":
    main()
